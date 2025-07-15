from langchain_community.document_loaders import PyPDFLoader, PDFPlumberLoader, UnstructuredPDFLoader, PyMuPDFLoader
from langchain_community.document_loaders import UnstructuredPowerPointLoader
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document
import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import numpy as np
import cv2
from pdf2image import convert_from_path
import tempfile
import shutil
import platform

class DocumentLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.extension = os.path.splitext(file_path)[1].lower()

    def load(self):
        """Load documents with enhanced PDF processing for scanned documents"""
        try:
            if self.extension == ".pdf":
                return self._load_pdf_with_ocr()
            elif self.extension == ".pptx":
                return UnstructuredPowerPointLoader(self.file_path).load()
            elif self.extension == ".docx":
                return UnstructuredWordDocumentLoader(self.file_path).load()
            elif self.extension == ".txt":
                return TextLoader(self.file_path).load()
            else:
                raise ValueError(f"Unsupported file type: {self.extension}")
        except Exception as e:
            print(f"[ERROR] Document loading failed for {self.file_path}: {e}")
            # Return a basic error document
            return [Document(
                page_content=f"Error loading document: {str(e)}. Please ensure the file is not corrupted and is in a supported format.",
                metadata={"page": 1, "source": self.file_path, "error": str(e)}
            )]

    def _load_pdf_with_ocr(self):
        """Enhanced PDF loading with OCR support for scanned documents"""
        try:
            # First, try to extract text using PyMuPDF (most reliable for text-based PDFs)
            print(f"[INFO] Attempting to extract text using PyMuPDF...")
            documents = self._extract_text_with_pymupdf()
            
            # Check if we got meaningful text content
            total_text = " ".join([doc.page_content for doc in documents])
            if len(total_text.strip()) > 50:  # If we have substantial text, use it
                print(f"[INFO] Successfully extracted {len(total_text)} characters using PyMuPDF")
                return documents
            
            # If text extraction failed or returned minimal content, try OCR
            print(f"[INFO] Text extraction returned minimal content ({len(total_text)} chars). Attempting OCR...")
            documents = self._extract_text_with_ocr()
            
            if documents:
                total_text = " ".join([doc.page_content for doc in documents])
                print(f"[INFO] Successfully extracted {len(total_text)} characters using OCR")
                return documents
            
            # If OCR also fails, try other PDF loaders as fallback
            print(f"[INFO] OCR failed. Trying alternative PDF loaders...")
            documents = self._try_alternative_pdf_loaders()
            
            if documents:
                total_text = " ".join([doc.page_content for doc in documents])
                print(f"[INFO] Successfully extracted {len(total_text)} characters using alternative loaders")
                return documents
            
            # If all methods fail, create a placeholder document with instructions
            print(f"[WARNING] All text extraction methods failed. Creating placeholder document.")
            return [Document(
                page_content="This appears to be a scanned document or image-based PDF. To enable full text extraction, please install Tesseract OCR. For now, you can still use the document for basic operations.",
                metadata={"page": 1, "source": self.file_path, "method": "placeholder"}
            )]
            
        except Exception as e:
            print(f"[ERROR] PDF processing failed: {e}")
            # Final fallback to basic PDF loader
            return PyPDFLoader(self.file_path).load()

    def _extract_text_with_pymupdf(self):
        """Extract text using PyMuPDF (handles most PDF types well)"""
        try:
            doc = fitz.open(self.file_path)
            documents = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try to extract text
                text = page.get_text()
                
                # If text is empty or very short, try to get text with more options
                if not text or len(text.strip()) < 10:
                    text = page.get_text("text")
                
                # If still no text, try to get text with layout preservation
                if not text or len(text.strip()) < 10:
                    text = page.get_text("dict")
                    # Extract text from the dict structure
                    if "blocks" in text:
                        text_content = []
                        for block in text["blocks"]:
                            if "lines" in block:
                                for line in block["lines"]:
                                    for span in line["spans"]:
                                        text_content.append(span["text"])
                        text = " ".join(text_content)
                
                if text and len(text.strip()) > 0:
                    documents.append(Document(
                        page_content=text.strip(),
                        metadata={"page": page_num + 1, "source": self.file_path}
                    ))
            
            doc.close()
            return documents
            
        except Exception as e:
            print(f"[WARNING] PyMuPDF extraction failed: {e}")
            return []

    def _extract_text_with_ocr(self):
        """Extract text from scanned PDFs using OCR"""
        try:
            # Check if Tesseract is available and configure it
            try:
                import pytesseract
                
                # Set Tesseract executable path explicitly
                tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                if os.path.exists(tesseract_path):
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path
                    print(f"[INFO] Tesseract found at: {tesseract_path}")
                else:
                    # Try to find tesseract in PATH
                    import subprocess
                    try:
                        result = subprocess.run(['tesseract', '--version'], capture_output=True, text=True)
                        if result.returncode == 0:
                            print("[INFO] Tesseract found in PATH")
                        else:
                            raise Exception("Tesseract not found in PATH")
                    except:
                        raise Exception("Tesseract executable not found")
                
                # Test if tesseract is working
                version = pytesseract.get_tesseract_version()
                print(f"[INFO] Tesseract version: {version}")
                
            except Exception as e:
                print(f"[WARNING] Tesseract not available: {e}")
                print("[INFO] Skipping OCR - Tesseract needs to be installed for OCR functionality")
                return []
            
            # Convert PDF to images
            print(f"[INFO] Converting PDF to images for OCR...")
            print(f"[DEBUG] pdftoppm path: {shutil.which('pdftoppm')}")
            print(f"[DEBUG] pdfinfo path: {shutil.which('pdfinfo')}")
            if platform.system() == "Windows":
                poppler_path = r"C:\\poppler\\poppler-23.11.0\\Library\\bin"
                images = convert_from_path(self.file_path, dpi=300, poppler_path=poppler_path)
            else:
                images = convert_from_path(self.file_path, dpi=300)
            
            documents = []
            
            for page_num, image in enumerate(images):
                print(f"[INFO] Processing page {page_num + 1} with OCR...")
                
                # Convert PIL image to OpenCV format for preprocessing
                img_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                
                # Preprocess image for better OCR (returns multiple versions)
                processed_images = self._preprocess_image_for_ocr(img_cv)
                
                # Convert all processed images to PIL format
                pil_images = []
                for processed_img in processed_images:
                    try:
                        pil_img = Image.fromarray(processed_img)
                        pil_images.append(pil_img)
                    except:
                        # If conversion fails, use original image
                        pil_images.append(image)
                
                # Perform OCR with multiple attempts and configurations
                best_text = ""
                best_length = 0
                
                # OCR configurations to try (in order of preference)
                ocr_configs = [
                    # Default configuration
                    {"config": "--oem 3 --psm 6", "name": "default"},
                    # Single uniform block of text
                    {"config": "--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?;:()[]{}'\"- ", "name": "alphanumeric"},
                    # Sparse text with OSD
                    {"config": "--oem 3 --psm 3", "name": "sparse_text"},
                    # Single text line
                    {"config": "--oem 3 --psm 7", "name": "single_line"},
                    # Single word
                    {"config": "--oem 3 --psm 8", "name": "single_word"},
                    # Single word in a circle
                    {"config": "--oem 3 --psm 9", "name": "circular_text"},
                    # Single character
                    {"config": "--oem 3 --psm 10", "name": "single_char"},
                    # Sparse text
                    {"config": "--oem 3 --psm 11", "name": "sparse_text_alt"},
                    # Raw line
                    {"config": "--oem 3 --psm 12", "name": "raw_line"},
                    # Uniform block of text
                    {"config": "--oem 3 --psm 13", "name": "uniform_block"}
                ]
                
                try:
                    # Try OCR on all preprocessed images with all configurations
                    for img_idx, pil_image in enumerate(pil_images):
                        for config in ocr_configs:
                            try:
                                text = pytesseract.image_to_string(
                                    pil_image, 
                                    config=config["config"],
                                    lang='eng'  # Specify English language
                                )
                                
                                # Clean the text
                                cleaned_text = self._clean_ocr_text(text)
                                
                                # Check if this configuration produced better results
                                if len(cleaned_text.strip()) > best_length:
                                    best_text = cleaned_text
                                    best_length = len(cleaned_text.strip())
                                    print(f"[INFO] Better OCR result with image {img_idx+1}, config {config['name']}: {best_length} characters")
                                
                            except Exception as config_error:
                                print(f"[DEBUG] OCR config {config['name']} failed for image {img_idx+1}: {config_error}")
                                continue
                    
                    # Use the best result
                    if best_text and len(best_text.strip()) > 10:
                        documents.append(Document(
                            page_content=best_text.strip(),
                            metadata={"page": page_num + 1, "source": self.file_path, "method": "OCR"}
                        ))
                        print(f"[INFO] OCR extracted {len(best_text)} characters from page {page_num + 1}")
                    else:
                        print(f"[WARNING] OCR returned minimal text for page {page_num + 1} (best: {best_length} chars)")
                        
                except Exception as e:
                    print(f"[WARNING] OCR failed for page {page_num + 1}: {e}")
                    continue
            
            return documents
            
        except Exception as e:
            print(f"[ERROR] OCR processing failed: {e}")
            return []

    def _preprocess_image_for_ocr(self, image):
        """Preprocess image for better OCR results"""
        try:
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply noise reduction
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Try multiple preprocessing approaches
            processed_images = []
            
            # Approach 1: Adaptive thresholding
            try:
                thresh1 = cv2.adaptiveThreshold(
                    denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
                )
                processed_images.append(thresh1)
            except:
                pass
            
            # Approach 2: Otsu thresholding
            try:
                _, thresh2 = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                processed_images.append(thresh2)
            except:
                pass
            
            # Approach 3: Simple thresholding
            try:
                _, thresh3 = cv2.threshold(denoised, 127, 255, cv2.THRESH_BINARY)
                processed_images.append(thresh3)
            except:
                pass
            
            # Approach 4: Original grayscale (sometimes works better)
            processed_images.append(denoised)
            
            # Apply morphological operations to clean up
            cleaned_images = []
            for img in processed_images:
                try:
                    # Small kernel for fine details
                    kernel_small = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
                    cleaned_small = cv2.morphologyEx(img, cv2.MORPH_CLOSE, kernel_small)
                    cleaned_images.append(cleaned_small)
                    
                    # Medium kernel for general cleaning
                    kernel_medium = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
                    cleaned_medium = cv2.morphologyEx(img, cv2.MORPH_CLOSE, kernel_medium)
                    cleaned_images.append(cleaned_medium)
                    
                except:
                    cleaned_images.append(img)
            
            # Return all processed images for testing
            return cleaned_images
            
        except Exception as e:
            print(f"[WARNING] Image preprocessing failed: {e}")
            return [image]

    def _clean_ocr_text(self, text):
        """Clean and improve OCR text"""
        if not text:
            return text
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Fix common OCR errors
        text = text.replace('|', 'I')  # Common OCR error
        text = text.replace('0', 'O')  # Sometimes numbers are confused with letters
        text = text.replace('1', 'l')  # Sometimes 1 is confused with l
        text = text.replace('l', 'I')  # Sometimes l is confused with I
        text = text.replace('rn', 'm')  # Common OCR error
        text = text.replace('cl', 'd')  # Common OCR error
        text = text.replace('vv', 'w')  # Common OCR error
        
        # Remove lines that are likely noise (very short lines)
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            # Keep lines with more than 2 characters and not just punctuation
            if len(line) > 2 and not all(c in '.,!?;:()[]{}' for c in line):
                cleaned_lines.append(line)
        
        # Join lines and clean up
        result = '\n'.join(cleaned_lines)
        
        # Remove excessive newlines
        result = '\n'.join(line for line in result.split('\n') if line.strip())
        
        return result

    def _try_alternative_pdf_loaders(self):
        """Try alternative PDF loaders if primary methods fail"""
        loaders = [
            ("PDFPlumberLoader", lambda: PDFPlumberLoader(self.file_path).load()),
            ("UnstructuredPDFLoader", lambda: UnstructuredPDFLoader(self.file_path).load()),
            ("PyPDFLoader", lambda: PyPDFLoader(self.file_path).load())
        ]
        
        for loader_name, loader_func in loaders:
            try:
                print(f"[INFO] Trying {loader_name}...")
                documents = loader_func()
                total_text = " ".join([doc.page_content for doc in documents])
                if len(total_text.strip()) > 10:
                    print(f"[INFO] {loader_name} successfully extracted {len(total_text)} characters")
                    return documents
            except Exception as e:
                print(f"[WARNING] {loader_name} failed: {e}")
                continue
        
        return []

    def get_page_count(self):
        """Get page count for different document types"""
        if self.extension == ".pdf":
            try:
                # Try PyMuPDF first (most reliable)
                doc = fitz.open(self.file_path)
                page_count = len(doc)
                doc.close()
                return page_count
            except Exception:
                try:
                    # Fallback to PyPDF2
                    import PyPDF2
                    with open(self.file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        return len(reader.pages)
                except Exception:
                    return None
        elif self.extension == ".pptx":
            try:
                from pptx import Presentation
                prs = Presentation(self.file_path)
                return len(prs.slides)
            except Exception:
                return None
        elif self.extension == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(self.file_path)
                # DOCX doesn't have strict pages, but we can estimate by section breaks or paragraphs
                return max(1, len(doc.paragraphs) // 30)  # Rough estimate: 30 paragraphs per page
            except Exception:
                return None
        elif self.extension == ".txt":
            try:
                with open(self.file_path, "r", encoding="utf-8") as f:
                    words = f.read().split()
                    return max(1, len(words) // 500)
            except Exception:
                return None
        else:
            return None 